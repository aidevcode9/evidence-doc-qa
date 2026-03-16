# app/services/export_service.py
"""Q&A Export Service for PDF and DOCX generation (FR-032)."""

from __future__ import annotations

import io
import json
from datetime import datetime, timezone
from typing import Any

from reportlab.lib import colors  # type: ignore[import-untyped]
from reportlab.lib.pagesizes import letter  # type: ignore[import-untyped]
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet  # type: ignore[import-untyped]
from reportlab.lib.units import inch  # type: ignore[import-untyped]
from reportlab.platypus import (  # type: ignore[import-untyped]
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from docx import Document as DocxDocument
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.db import QAMessage, QASession


def generate_pdf_export(
    session: QASession,
    messages: list[QAMessage],
) -> bytes:
    """Generate PDF with Q&A history and citations."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )

    styles = getSampleStyleSheet()

    # Custom styles
    title_style = ParagraphStyle(
        "TitleStyle",
        parent=styles["Title"],
        fontSize=18,
        spaceAfter=6,
    )

    header_style = ParagraphStyle(
        "HeaderStyle",
        parent=styles["Normal"],
        fontSize=10,
        textColor=colors.gray,
        spaceAfter=12,
    )

    question_style = ParagraphStyle(
        "QuestionStyle",
        parent=styles["Normal"],
        fontSize=11,
        fontName="Helvetica-Bold",
        textColor=colors.HexColor("#1a365d"),
        spaceBefore=12,
        spaceAfter=6,
    )

    answer_style = ParagraphStyle(
        "AnswerStyle",
        parent=styles["Normal"],
        fontSize=11,
        spaceBefore=6,
        spaceAfter=6,
        leftIndent=12,
    )

    citation_header_style = ParagraphStyle(
        "CitationHeaderStyle",
        parent=styles["Normal"],
        fontSize=10,
        fontName="Helvetica-Bold",
        textColor=colors.HexColor("#2d3748"),
        spaceBefore=8,
        spaceAfter=4,
    )

    citation_style = ParagraphStyle(
        "CitationStyle",
        parent=styles["Normal"],
        fontSize=9,
        textColor=colors.HexColor("#4a5568"),
        leftIndent=24,
        spaceAfter=4,
    )

    refusal_style = ParagraphStyle(
        "RefusalStyle",
        parent=styles["Normal"],
        fontSize=10,
        textColor=colors.HexColor("#c53030"),
        fontName="Helvetica-Oblique",
        leftIndent=12,
        spaceAfter=6,
    )

    footer_style = ParagraphStyle(
        "FooterStyle",
        parent=styles["Normal"],
        fontSize=8,
        textColor=colors.gray,
        spaceBefore=20,
    )

    elements: list[Any] = []

    # Title
    elements.append(Paragraph("EVIDENCE BOUND - Q&A Export", title_style))

    # Header info
    generated_at = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")
    session_short = session.session_id[:8] if len(session.session_id) > 8 else session.session_id
    header_text = f"Session: {session_short} | Generated: {generated_at}"
    elements.append(Paragraph(header_text, header_style))
    elements.append(Spacer(1, 12))

    # Horizontal line
    line_data = [["" * 80]]
    line_table = Table(line_data, colWidths=[7 * inch])
    line_table.setStyle(
        TableStyle([("LINEBELOW", (0, 0), (-1, -1), 1, colors.HexColor("#e2e8f0"))])
    )
    elements.append(line_table)
    elements.append(Spacer(1, 12))

    # Process messages in pairs (user question + assistant response)
    i = 0
    while i < len(messages):
        msg = messages[i]

        if msg.role == "user":
            # Question
            q_text = _escape_xml(msg.content)
            elements.append(Paragraph(f"Q: {q_text}", question_style))

            # Look for following assistant message
            if i + 1 < len(messages) and messages[i + 1].role == "assistant":
                assistant_msg = messages[i + 1]
                a_text = _escape_xml(assistant_msg.content)
                elements.append(Paragraph(f"A: {a_text}", answer_style))

                # Check for refusal
                if assistant_msg.refusal_code:
                    refusal_text = f"[Refusal: {assistant_msg.refusal_code}]"
                    elements.append(Paragraph(refusal_text, refusal_style))

                # Citations
                if assistant_msg.citations_json:
                    try:
                        citations = json.loads(assistant_msg.citations_json)
                        if citations:
                            elements.append(Paragraph("CITATIONS:", citation_header_style))
                            for cit in citations:
                                idx = cit.get("citation_index", "?")
                                doc_name = cit.get("doc_name", "Unknown")
                                page = cit.get("page_num", "?")
                                snippet = _escape_xml(cit.get("snippet", "")[:200])
                                cit_text = f"[{idx}] {doc_name}, Page {page}<br/>&quot;{snippet}...&quot;"
                                elements.append(Paragraph(cit_text, citation_style))
                    except json.JSONDecodeError:
                        pass

                # Evidence grade
                if assistant_msg.evidence_json:
                    try:
                        evidence = json.loads(assistant_msg.evidence_json)
                        grade = evidence.get("evidence_grade", "")
                        label = evidence.get("evidence_label", "")
                        if grade:
                            grade_text = f"Evidence Grade: {grade} ({label})"
                            elements.append(Paragraph(grade_text, citation_style))
                    except json.JSONDecodeError:
                        pass

                i += 2
            else:
                i += 1

            elements.append(Spacer(1, 8))
            # Separator line
            elements.append(line_table)
            elements.append(Spacer(1, 8))

        else:
            # Standalone assistant message (shouldn't happen normally)
            a_text = _escape_xml(msg.content)
            elements.append(Paragraph(f"A: {a_text}", answer_style))
            i += 1

    # Footer
    elements.append(Spacer(1, 20))
    footer_text = f"Docs Snapshot: {session.docs_snapshot_id}"
    elements.append(Paragraph(footer_text, footer_style))

    doc.build(elements)
    return buffer.getvalue()


def generate_docx_export(
    session: QASession,
    messages: list[QAMessage],
) -> bytes:
    """Generate DOCX with Q&A history and citations."""
    document = DocxDocument()

    # Title
    title = document.add_heading("EVIDENCE BOUND - Q&A Export", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Header info
    generated_at = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")
    session_short = session.session_id[:8] if len(session.session_id) > 8 else session.session_id
    header_para = document.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header_para.add_run(f"Session: {session_short} | Generated: {generated_at}")
    header_run.font.size = Pt(10)
    header_run.font.italic = True

    document.add_paragraph()  # Spacer

    # Process messages in pairs
    i = 0
    while i < len(messages):
        msg = messages[i]

        if msg.role == "user":
            # Question
            q_para = document.add_paragraph()
            q_run = q_para.add_run(f"Q: {msg.content}")
            q_run.font.bold = True
            q_run.font.size = Pt(11)

            # Look for following assistant message
            if i + 1 < len(messages) and messages[i + 1].role == "assistant":
                assistant_msg = messages[i + 1]

                # Answer
                a_para = document.add_paragraph()
                a_para.paragraph_format.left_indent = Inches(0.25)
                a_run = a_para.add_run(f"A: {assistant_msg.content}")
                a_run.font.size = Pt(11)

                # Check for refusal
                if assistant_msg.refusal_code:
                    refusal_para = document.add_paragraph()
                    refusal_para.paragraph_format.left_indent = Inches(0.25)
                    refusal_run = refusal_para.add_run(f"[Refusal: {assistant_msg.refusal_code}]")
                    refusal_run.font.italic = True
                    refusal_run.font.size = Pt(10)

                # Citations
                if assistant_msg.citations_json:
                    try:
                        citations = json.loads(assistant_msg.citations_json)
                        if citations:
                            cit_header = document.add_paragraph()
                            cit_header.paragraph_format.left_indent = Inches(0.25)
                            cit_header_run = cit_header.add_run("CITATIONS:")
                            cit_header_run.font.bold = True
                            cit_header_run.font.size = Pt(10)

                            for cit in citations:
                                idx = cit.get("citation_index", "?")
                                doc_name = cit.get("doc_name", "Unknown")
                                page = cit.get("page_num", "?")
                                snippet = cit.get("snippet", "")[:200]

                                cit_para = document.add_paragraph()
                                cit_para.paragraph_format.left_indent = Inches(0.5)

                                cit_ref = cit_para.add_run(f"[{idx}] {doc_name}, Page {page}")
                                cit_ref.font.bold = True
                                cit_ref.font.size = Pt(9)

                                cit_para.add_run("\n")
                                cit_quote = cit_para.add_run(f'"{snippet}..."')
                                cit_quote.font.italic = True
                                cit_quote.font.size = Pt(9)
                    except json.JSONDecodeError:
                        pass

                # Evidence grade
                if assistant_msg.evidence_json:
                    try:
                        evidence = json.loads(assistant_msg.evidence_json)
                        grade = evidence.get("evidence_grade", "")
                        label = evidence.get("evidence_label", "")
                        if grade:
                            grade_para = document.add_paragraph()
                            grade_para.paragraph_format.left_indent = Inches(0.25)
                            grade_run = grade_para.add_run(f"Evidence Grade: {grade} ({label})")
                            grade_run.font.size = Pt(9)
                    except json.JSONDecodeError:
                        pass

                i += 2
            else:
                i += 1

            # Separator
            document.add_paragraph("─" * 60)

        else:
            # Standalone assistant message
            a_para = document.add_paragraph()
            a_run = a_para.add_run(f"A: {msg.content}")
            a_run.font.size = Pt(11)
            i += 1

    # Footer
    document.add_paragraph()
    footer_para = document.add_paragraph()
    footer_run = footer_para.add_run(f"Docs Snapshot: {session.docs_snapshot_id}")
    footer_run.font.size = Pt(8)
    footer_run.font.italic = True

    # Save to bytes
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


###############################################################################
# FR-033: Cited-Only Packet Export
###############################################################################

ExhibitDict = dict[str, Any]


def extract_cited_exhibits(messages: list[QAMessage]) -> list[ExhibitDict]:
    """Extract deduplicated exhibits from session citations.

    Groups citations by document, collecting unique pages and snippets.
    Returns a list of exhibit dicts sorted by doc_name.
    """
    # doc_id → exhibit info
    exhibits: dict[str, ExhibitDict] = {}

    for msg in messages:
        if msg.role != "assistant" or not msg.citations_json:
            continue
        try:
            citations = json.loads(msg.citations_json)
        except (json.JSONDecodeError, TypeError):
            continue
        if not isinstance(citations, list):
            continue

        for cit in citations:
            doc_id = cit.get("doc_id", "unknown")
            if doc_id not in exhibits:
                exhibits[doc_id] = {
                    "doc_id": doc_id,
                    "doc_name": cit.get("doc_name", "Unknown Document"),
                    "pages": set(),
                    "snippets": [],
                    "_seen_chunks": set(),
                }

            # Collect pages
            page = cit.get("page_num")
            if page is not None:
                exhibits[doc_id]["pages"].add(page)
            page_end = cit.get("page_end")
            if page_end is not None and page_end != page:
                exhibits[doc_id]["pages"].add(page_end)

            # Collect unique snippets (deduplicate by chunk_id)
            chunk_id = cit.get("chunk_id", "")
            if chunk_id and chunk_id not in exhibits[doc_id]["_seen_chunks"]:
                exhibits[doc_id]["_seen_chunks"].add(chunk_id)
                exhibits[doc_id]["snippets"].append({
                    "text": cit.get("snippet", ""),
                    "page_num": page or 0,
                    "chunk_id": chunk_id,
                })

    # Convert sets to sorted lists, remove internal tracking
    result: list[ExhibitDict] = []
    for exhibit in sorted(exhibits.values(), key=lambda e: e["doc_name"]):
        result.append({
            "doc_id": exhibit["doc_id"],
            "doc_name": exhibit["doc_name"],
            "pages": sorted(exhibit["pages"]),
            "snippets": exhibit["snippets"],
        })
    return result


def generate_cited_packet_pdf(
    session: QASession,
    messages: list[QAMessage],
) -> bytes:
    """Generate PDF cited-only packet listing referenced exhibits/pages."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )

    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "CPTitle", parent=styles["Title"], fontSize=18, spaceAfter=6,
    )
    header_style = ParagraphStyle(
        "CPHeader", parent=styles["Normal"], fontSize=10,
        textColor=colors.gray, spaceAfter=12,
    )
    exhibit_title_style = ParagraphStyle(
        "CPExhibitTitle", parent=styles["Normal"], fontSize=12,
        fontName="Helvetica-Bold", textColor=colors.HexColor("#1a365d"),
        spaceBefore=14, spaceAfter=4,
    )
    page_style = ParagraphStyle(
        "CPPages", parent=styles["Normal"], fontSize=10,
        textColor=colors.HexColor("#2d3748"), leftIndent=12, spaceAfter=4,
    )
    snippet_style = ParagraphStyle(
        "CPSnippet", parent=styles["Normal"], fontSize=9,
        textColor=colors.HexColor("#4a5568"), leftIndent=24, spaceAfter=4,
    )
    footer_style = ParagraphStyle(
        "CPFooter", parent=styles["Normal"], fontSize=8,
        textColor=colors.gray, spaceBefore=20,
    )
    empty_style = ParagraphStyle(
        "CPEmpty", parent=styles["Normal"], fontSize=11,
        textColor=colors.HexColor("#718096"), spaceBefore=20,
    )

    elements: list[Any] = []

    # Title
    elements.append(Paragraph("EVIDENCE BOUND - Cited Exhibits Packet", title_style))

    # Header
    generated_at = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")
    session_short = session.session_id[:8] if len(session.session_id) > 8 else session.session_id
    elements.append(Paragraph(
        f"Session: {session_short} | Generated: {generated_at}", header_style,
    ))
    elements.append(Spacer(1, 12))

    # Separator
    line_data = [["" * 80]]
    line_table = Table(line_data, colWidths=[7 * inch])
    line_table.setStyle(
        TableStyle([("LINEBELOW", (0, 0), (-1, -1), 1, colors.HexColor("#e2e8f0"))])
    )
    elements.append(line_table)
    elements.append(Spacer(1, 12))

    # Extract exhibits
    exhibits = extract_cited_exhibits(messages)

    if not exhibits:
        elements.append(Paragraph("No cited exhibits found in this session.", empty_style))
    else:
        # Summary table
        elements.append(Paragraph(
            f"<b>{len(exhibits)}</b> document(s) cited in this session:",
            page_style,
        ))
        elements.append(Spacer(1, 8))

        for idx, exhibit in enumerate(exhibits, start=1):
            pages_str = ", ".join(str(p) for p in exhibit["pages"])
            elements.append(Paragraph(
                f"Exhibit {idx}: {_escape_xml(exhibit['doc_name'])}",
                exhibit_title_style,
            ))
            elements.append(Paragraph(f"Pages cited: {pages_str}", page_style))

            for snippet in exhibit["snippets"]:
                text = _escape_xml(snippet["text"][:300])
                elements.append(Paragraph(
                    f"p.{snippet['page_num']}: &quot;{text}...&quot;",
                    snippet_style,
                ))

            elements.append(Spacer(1, 6))
            elements.append(line_table)
            elements.append(Spacer(1, 6))

    # Footer
    elements.append(Spacer(1, 20))
    elements.append(Paragraph(
        f"Docs Snapshot: {session.docs_snapshot_id}", footer_style,
    ))

    doc.build(elements)
    return buffer.getvalue()


def generate_cited_packet_docx(
    session: QASession,
    messages: list[QAMessage],
) -> bytes:
    """Generate DOCX cited-only packet listing referenced exhibits/pages."""
    document = DocxDocument()

    # Title
    title = document.add_heading("EVIDENCE BOUND - Cited Exhibits Packet", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Header
    generated_at = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")
    session_short = session.session_id[:8] if len(session.session_id) > 8 else session.session_id
    header_para = document.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header_para.add_run(f"Session: {session_short} | Generated: {generated_at}")
    header_run.font.size = Pt(10)
    header_run.font.italic = True

    document.add_paragraph()

    # Extract exhibits
    exhibits = extract_cited_exhibits(messages)

    if not exhibits:
        empty_para = document.add_paragraph()
        empty_run = empty_para.add_run("No cited exhibits found in this session.")
        empty_run.font.italic = True
        empty_run.font.size = Pt(11)
    else:
        # Summary
        summary = document.add_paragraph()
        summary_run = summary.add_run(f"{len(exhibits)} document(s) cited in this session:")
        summary_run.font.bold = True
        summary_run.font.size = Pt(11)

        document.add_paragraph()

        for idx, exhibit in enumerate(exhibits, start=1):
            pages_str = ", ".join(str(p) for p in exhibit["pages"])

            # Exhibit header
            ex_para = document.add_paragraph()
            ex_run = ex_para.add_run(f"Exhibit {idx}: {exhibit['doc_name']}")
            ex_run.font.bold = True
            ex_run.font.size = Pt(12)

            # Pages
            pg_para = document.add_paragraph()
            pg_para.paragraph_format.left_indent = Inches(0.25)
            pg_run = pg_para.add_run(f"Pages cited: {pages_str}")
            pg_run.font.size = Pt(10)

            # Snippets
            for snippet in exhibit["snippets"]:
                sn_para = document.add_paragraph()
                sn_para.paragraph_format.left_indent = Inches(0.5)
                sn_ref = sn_para.add_run(f"p.{snippet['page_num']}: ")
                sn_ref.font.bold = True
                sn_ref.font.size = Pt(9)
                sn_quote = sn_para.add_run(f'"{snippet["text"][:300]}..."')
                sn_quote.font.italic = True
                sn_quote.font.size = Pt(9)

            document.add_paragraph("─" * 60)

    # Footer
    document.add_paragraph()
    footer_para = document.add_paragraph()
    footer_run = footer_para.add_run(f"Docs Snapshot: {session.docs_snapshot_id}")
    footer_run.font.size = Pt(8)
    footer_run.font.italic = True

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _escape_xml(text: str) -> str:
    """Escape special XML characters for PDF generation."""
    if not text:
        return ""
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
        .replace("'", "&apos;")
    )
